"""Document processor for extracting text and images from various document formats."""

import os
import io
import base64
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass
from PIL import Image
import fitz  # PyMuPDF
import pytesseract
import cv2
import numpy as np
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import config
from hierarchical_title_extractor import HierarchicalTitleExtractor

@dataclass
class DocumentChunk:
    """Represents a chunk of document content with associated metadata."""
    text: str
    page_number: int
    chunk_id: str
    document_name: str
    chunk_type: str  # 'text', 'image_caption', 'table'
    section_title: str = ""  # 章節標題
    
@dataclass
class ImageData:
    """Represents an extracted image with metadata and context."""
    image: Image.Image
    image_id: str
    page_number: int
    document_name: str
    caption: str
    bbox: Optional[Tuple[float, float, float, float]] = None  # x0, y0, x1, y1
    ocr_text: str = ""
    # 新增上下文資訊
    surrounding_text: str = ""      # 圖片前後的段落文字
    section_title: str = ""         # 所在章節標題（現在包含完整的編號格式）
    page_context: str = ""          # 整頁的摘要文字
    preceding_text: str = ""        # 圖片前面的文字
    following_text: str = ""        # 圖片後面的文字

class DocumentProcessor:
    """Processes documents to extract text chunks and images."""
    
    def __init__(self):
        self.supported_formats = config.SUPPORTED_FORMATS
        self.image_formats = config.IMAGE_FORMATS
        self.max_image_size = config.MAX_IMAGE_SIZE
        self.ocr_enabled = config.OCR_ENABLED
        self.title_extractor = HierarchicalTitleExtractor()
        
    def process_document(self, file_path: str, chunking_strategy: str = "page") -> Tuple[List[DocumentChunk], List[ImageData]]:
        """
        Process a document and extract text chunks and images.
        
        Args:
            file_path: Path to the document file
            chunking_strategy: "page" for page-based chunking, "section" for section-based chunking
            
        Returns:
            Tuple of (text_chunks, images)
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        document_name = os.path.basename(file_path)
        
        if file_ext == '.pdf':
            return self._process_pdf(file_path, document_name, chunking_strategy)
        elif file_ext in ['.docx', '.doc']:
            return self._process_docx(file_path, document_name)
        elif file_ext == '.txt':
            return self._process_txt(file_path, document_name)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _process_pdf(self, file_path: str, document_name: str, 
                    chunking_strategy: str = "page") -> Tuple[List[DocumentChunk], List[ImageData]]:
        """
        Process PDF document.
        
        Args:
            file_path: Path to PDF file
            document_name: Name of the document
            chunking_strategy: "page" for page-based chunking, "section" for section-based chunking
        """
        text_chunks = []
        images = []
        
        doc = fitz.open(file_path)
        
        # First, build the section index for the entire document
        print("🔍 Building section index for document...")
        self.title_extractor.build_document_section_index(doc)
        all_sections = self.title_extractor.get_all_sections()
        print(f"🔍 DEBUG: Found {len(all_sections)} sections: {all_sections[:5]}{'...' if len(all_sections) > 5 else ''}")
        
        # Get full document text for position calculation
        full_document_text = ""
        page_start_positions = {}
        current_pos = 0
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text()
            page_start_positions[page_num] = current_pos
            full_document_text += page_text + "\n"
            current_pos += len(page_text) + 1
        
        # Choose chunking strategy
        if chunking_strategy == "section":
            print("🔧 Using section-based chunking strategy")
            text_chunks = self._split_text_into_chunks_by_sections(
                full_document_text, document_name, doc
            )
        else:
            print("🔧 Using page-based chunking strategy")
            # Maintain section continuity across pages (original method)
            last_chunk_sections = []
            overall_chunk_index = 0
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Extract text
                text = page.get_text()
                if text.strip():
                    # Split text into smaller chunks with enhanced section information
                    chunks, last_chunk_sections = self._split_text_into_chunks_with_sections(
                        text, page_num + 1, document_name, 
                        page_start_positions[page_num], last_chunk_sections, overall_chunk_index
                    )
                    text_chunks.extend(chunks)
                    overall_chunk_index += len(chunks)
                    
                    print(f"📄 頁數{page_num + 1}: 分塊數={len(chunks)}")
        
        # Extract images with context (independent of chunking strategy)
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()  # Get page text for context
            
            image_list = page.get_images()
            for img_index, img in enumerate(image_list):
                try:
                    # Get image data
                    xref = img[0]
                    pix = fitz.Pixmap(doc, xref)
                    
                    if pix.n - pix.alpha < 4:  # Skip if not RGB/GRAY
                        img_data = pix.tobytes("png")
                        pil_image = Image.open(io.BytesIO(img_data))
                        
                        # Resize if too large
                        pil_image = self._resize_image(pil_image)
                        
                        # Extract OCR text if enabled
                        ocr_text = ""
                        if self.ocr_enabled:
                            ocr_text = self._extract_ocr_text(pil_image)
                        
                        # Try to find caption near the image
                        caption = self._find_image_caption(page, img, page_num + 1)
                        
                        # Extract image context using PyMuPDF built-in functions
                        context_info = self._extract_image_context(page, img, text, page_num + 1, img_index)
                        
                        image_data = ImageData(
                            image=pil_image,
                            image_id=f"{document_name}_page{page_num+1}_img{img_index}",
                            page_number=page_num + 1,
                            document_name=document_name,
                            caption=caption,
                            bbox=context_info.get('bbox'),  # Get bbox from context_info if available
                            ocr_text=ocr_text,
                            # 新增的上下文資訊
                            surrounding_text=context_info['surrounding_text'],
                            section_title=context_info['section_title'],
                            page_context=context_info['page_context'],
                            preceding_text=context_info['preceding_text'],
                            following_text=context_info['following_text']
                        )
                        images.append(image_data)
                        
                        print(f"📸 已提取圖片上下文: {document_name}_page{page_num+1}_img{img_index}")
                        print(f"   章節標題: {context_info['section_title'][:50]}...")
                        print(f"   周圍文字: {context_info['surrounding_text'][:100]}...")
                    
                    pix = None
                except Exception as e:
                    print(f"Error processing image {img_index} on page {page_num + 1}: {e}")
                    continue
        
        doc.close()
        return text_chunks, images
    
    def _process_docx(self, file_path: str, document_name: str) -> Tuple[List[DocumentChunk], List[ImageData]]:
        """Process DOCX document."""
        text_chunks = []
        images = []
        
        doc = Document(file_path)
        page_num = 1  # DOCX doesn't have explicit page numbers
        
        # Extract potential section titles from document
        section_title = self._extract_docx_section_title(doc)
        
        # Process paragraphs and extract embedded images
        for element in doc.element.body:
            if isinstance(element, CT_P):
                paragraph = Paragraph(element, doc)
                text = paragraph.text.strip()
                
                if text:
                    chunk = DocumentChunk(
                        text=text,
                        page_number=page_num,
                        chunk_id=f"{document_name}_p{len(text_chunks)}",
                        document_name=document_name,
                        chunk_type="text",
                        section_title=section_title
                    )
                    text_chunks.append(chunk)
                
                # Check for images in paragraph
                for run in paragraph.runs:
                    for inline_shape in run._element.xpath('.//a:blip'):
                        try:
                            # Extract embedded image
                            image_part = doc.part.related_parts[inline_shape.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')]
                            image_data = image_part.blob
                            pil_image = Image.open(io.BytesIO(image_data))
                            pil_image = self._resize_image(pil_image)
                            
                            # Extract OCR text if enabled
                            ocr_text = ""
                            if self.ocr_enabled:
                                ocr_text = self._extract_ocr_text(pil_image)
                            
                            image_obj = ImageData(
                                image=pil_image,
                                image_id=f"{document_name}_img{len(images)}",
                                page_number=page_num,
                                document_name=document_name,
                                caption=text if text else "",
                                ocr_text=ocr_text
                            )
                            images.append(image_obj)
                        except Exception as e:
                            print(f"Error extracting image from DOCX: {e}")
                            continue
            
            elif isinstance(element, CT_Tbl):
                table = Table(element, doc)
                table_text = self._extract_table_text(table)
                if table_text:
                    chunk = DocumentChunk(
                        text=table_text,
                        page_number=page_num,
                        chunk_id=f"{document_name}_table{len(text_chunks)}",
                        document_name=document_name,
                        chunk_type="table",
                        section_title=section_title
                    )
                    text_chunks.append(chunk)
        
        return text_chunks, images
    
    def _process_txt(self, file_path: str, document_name: str) -> Tuple[List[DocumentChunk], List[ImageData]]:
        """Process plain text document."""
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        # Extract section title from first few lines of text
        section_title = self._extract_txt_section_title(text)
        
        chunks = self._split_text_into_chunks(text, 1, document_name, section_title)
        return chunks, []  # No images in plain text
    
    def _split_text_into_chunks(self, text: str, page_number: int, document_name: str, 
                               section_title: str = "", chunk_size: int = 500) -> List[DocumentChunk]:
        """Split text into smaller chunks for better retrieval."""
        chunks = []
        sentences = text.split('. ')
        current_chunk = ""
        chunk_index = 0
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) < chunk_size:
                current_chunk += sentence + ". "
            else:
                if current_chunk.strip():
                    chunk = DocumentChunk(
                        text=current_chunk.strip(),
                        page_number=page_number,
                        chunk_id=f"{document_name}_p{page_number}_c{chunk_index}",
                        document_name=document_name,
                        chunk_type="text",
                        section_title=section_title
                    )
                    chunks.append(chunk)
                    chunk_index += 1
                current_chunk = sentence + ". "
        
        # Add the last chunk
        if current_chunk.strip():
            chunk = DocumentChunk(
                text=current_chunk.strip(),
                page_number=page_number,
                chunk_id=f"{document_name}_p{page_number}_c{chunk_index}",
                document_name=document_name,
                chunk_type="text",
                section_title=section_title
            )
            chunks.append(chunk)
        
        return chunks
    
    def _split_text_into_chunks_with_sections(self, text: str, page_number: int, 
                                            document_name: str, page_start_pos: int,
                                            last_chunk_sections: List[str], 
                                            start_chunk_index: int,
                                            chunk_size: int = 500) -> tuple[List[DocumentChunk], List[str]]:
        """Split text into chunks with enhanced section information."""
        chunks = []
        sentences = text.split('. ')
        current_chunk = ""
        chunk_index = start_chunk_index
        current_pos = page_start_pos
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) < chunk_size:
                current_chunk += sentence + ". "
            else:
                if current_chunk.strip():
                    # Analyze sections for this chunk using optimized method
                    print(f"🔍 DEBUG: Calling analyze_chunk_sections with last_chunk_sections: {last_chunk_sections}")
                    chunk_sections = self.title_extractor.analyze_chunk_sections(
                        current_chunk, current_pos, last_chunk_sections
                    )
                    
                    chunk = DocumentChunk(
                        text=current_chunk.strip(),
                        page_number=page_number,
                        chunk_id=f"{document_name}_p{page_number}_c{chunk_index}",
                        document_name=document_name,
                        chunk_type="text",
                        section_title=""  # Keep for backward compatibility
                    )
                    
                    # Add the enhanced section information
                    chunk.all_sections = chunk_sections
                    last_chunk_sections = chunk_sections  # Update for next chunk
                    
                    chunks.append(chunk)
                    chunk_index += 1
                    
                current_chunk = sentence + ". "
                current_pos += len(sentence) + 2  # +2 for ". "
        
        # Add the last chunk
        if current_chunk.strip():
            chunk_sections = self.title_extractor.analyze_chunk_sections(
                current_chunk, current_pos, last_chunk_sections
            )
            
            chunk = DocumentChunk(
                text=current_chunk.strip(),
                page_number=page_number,
                chunk_id=f"{document_name}_p{page_number}_c{chunk_index}",
                document_name=document_name,
                chunk_type="text",
                section_title=""
            )
            chunk.all_sections = chunk_sections
            last_chunk_sections = chunk_sections
            chunks.append(chunk)
        
        return chunks, last_chunk_sections
    
    def _split_text_into_chunks_by_sections(self, full_document_text: str, document_name: str, 
                                          doc: fitz.Document, max_chunk_size: int = 1000) -> List[DocumentChunk]:
        """
        Split text into chunks based on section boundaries.
        
        Args:
            full_document_text: Complete document text
            document_name: Name of the document
            doc: PyMuPDF document object for page mapping
            max_chunk_size: Maximum size for each chunk
            
        Returns:
            List of DocumentChunk objects based on sections
        """
        chunks = []
        
        # Get all sections from the title extractor
        section_index = self.title_extractor.section_index
        all_sections = list(section_index.keys())
        
        print(f"🔧 Section-based chunking: Found {len(all_sections)} sections")
        
        if not all_sections:
            # Fallback: if no sections found, use simple text splitting
            print("⚠️ No sections found, falling back to simple text splitting")
            return self._simple_text_splitting(full_document_text, document_name, doc, max_chunk_size)
        
        # Sort sections by their appearance order in the document
        sections_with_positions = []
        for section_title in all_sections:
            section_info = section_index[section_title]
            page_num = section_info.get('page_number', 0)
            y_pos = section_info.get('y_position', 0)
            estimated_pos = page_num * 10000 + y_pos  # Rough position estimate
            sections_with_positions.append((estimated_pos, section_title, section_info))
        
        sections_with_positions.sort(key=lambda x: x[0])
        
        # Find section boundaries in the full text
        section_boundaries = []
        for i, (pos, section_title, section_info) in enumerate(sections_with_positions):
            # Try to find the section title in the full text
            section_pattern = section_info.get('full_title', section_title)
            
            # Search for the section title pattern
            section_start = full_document_text.find(section_pattern)
            if section_start == -1:
                # Try alternative patterns
                number = section_info.get('number', '')
                title = section_info.get('title', '')
                if number and title:
                    alt_pattern = f"{number} {title}"
                    section_start = full_document_text.find(alt_pattern)
            
            if section_start != -1:
                section_boundaries.append({
                    'start_pos': section_start,
                    'section_title': section_title,
                    'section_info': section_info,
                    'full_title': section_pattern
                })
                print(f"📍 Found section '{section_title}' at position {section_start}")
        
        # Sort boundaries by position
        section_boundaries.sort(key=lambda x: x['start_pos'])
        
        # Create chunks based on section boundaries
        chunk_index = 0
        for i, boundary in enumerate(section_boundaries):
            start_pos = boundary['start_pos']
            
            # Determine end position (start of next section or end of document)
            if i + 1 < len(section_boundaries):
                end_pos = section_boundaries[i + 1]['start_pos']
            else:
                end_pos = len(full_document_text)
            
            # Extract section text
            section_text = full_document_text[start_pos:end_pos].strip()
            
            if not section_text:
                continue
                
            # Determine which page this section starts on
            page_number = boundary['section_info'].get('page_number', 1)
            
            # If section is too large, split it further
            if len(section_text) <= max_chunk_size:
                # Single chunk for this section
                chunk = DocumentChunk(
                    text=section_text,
                    page_number=page_number,
                    chunk_id=f"{document_name}_section_{chunk_index}",
                    document_name=document_name,
                    chunk_type="section",
                    section_title=boundary['section_title']
                )
                # Set the primary section for this chunk
                chunk.all_sections = [boundary['section_title']]
                chunks.append(chunk)
                chunk_index += 1
                
                print(f"📦 Created section chunk: '{boundary['section_title']}' ({len(section_text)} chars)")
                
            else:
                # Split large section into smaller chunks
                sub_chunks = self._split_large_section(
                    section_text, 
                    boundary['section_title'], 
                    page_number, 
                    document_name, 
                    chunk_index, 
                    max_chunk_size
                )
                chunks.extend(sub_chunks)
                chunk_index += len(sub_chunks)
                
                print(f"📦 Split large section '{boundary['section_title']}' into {len(sub_chunks)} chunks")
        
        print(f"✅ Section-based chunking complete: {len(chunks)} total chunks")
        return chunks
    
    def _split_large_section(self, section_text: str, section_title: str, page_number: int,
                           document_name: str, start_chunk_index: int, max_chunk_size: int) -> List[DocumentChunk]:
        """Split a large section into smaller chunks."""
        chunks = []
        
        # Split by paragraphs first (double newlines)
        paragraphs = section_text.split('\n\n')
        
        current_chunk = ""
        chunk_index = start_chunk_index
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
                
            # If adding this paragraph would exceed max size, create a chunk
            if current_chunk and len(current_chunk) + len(paragraph) + 2 > max_chunk_size:
                chunk = DocumentChunk(
                    text=current_chunk.strip(),
                    page_number=page_number,
                    chunk_id=f"{document_name}_section_{chunk_index}",
                    document_name=document_name,
                    chunk_type="section_part",
                    section_title=section_title
                )
                chunk.all_sections = [section_title]
                chunks.append(chunk)
                chunk_index += 1
                current_chunk = ""
            
            current_chunk += paragraph + "\n\n"
        
        # Add remaining text as final chunk
        if current_chunk.strip():
            chunk = DocumentChunk(
                text=current_chunk.strip(),
                page_number=page_number,
                chunk_id=f"{document_name}_section_{chunk_index}",
                document_name=document_name,
                chunk_type="section_part",
                section_title=section_title
            )
            chunk.all_sections = [section_title]
            chunks.append(chunk)
        
        return chunks
    
    def _simple_text_splitting(self, text: str, document_name: str, doc: fitz.Document, 
                             max_chunk_size: int) -> List[DocumentChunk]:
        """Fallback method for simple text splitting when no sections are found."""
        chunks = []
        sentences = text.split('. ')
        current_chunk = ""
        chunk_index = 0
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) < max_chunk_size:
                current_chunk += sentence + ". "
            else:
                if current_chunk.strip():
                    chunk = DocumentChunk(
                        text=current_chunk.strip(),
                        page_number=1,  # Default page
                        chunk_id=f"{document_name}_simple_{chunk_index}",
                        document_name=document_name,
                        chunk_type="text",
                        section_title=""
                    )
                    chunk.all_sections = []
                    chunks.append(chunk)
                    chunk_index += 1
                current_chunk = sentence + ". "
        
        # Add final chunk
        if current_chunk.strip():
            chunk = DocumentChunk(
                text=current_chunk.strip(),
                page_number=1,
                chunk_id=f"{document_name}_simple_{chunk_index}",
                document_name=document_name,
                chunk_type="text",
                section_title=""
            )
            chunk.all_sections = []
            chunks.append(chunk)
        
        return chunks
    
    def _resize_image(self, image: Image.Image) -> Image.Image:
        """Resize image to maximum allowed size while maintaining aspect ratio."""
        if image.size[0] > self.max_image_size[0] or image.size[1] > self.max_image_size[1]:
            image.thumbnail(self.max_image_size, Image.Resampling.LANCZOS)
        return image
    
    def _extract_ocr_text(self, image: Image.Image) -> str:
        """Extract text from image using OCR."""
        try:
            # Convert PIL image to numpy array for OpenCV
            img_array = np.array(image)
            if len(img_array.shape) == 3:
                img_array = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
            
            # Use pytesseract for OCR
            ocr_text = pytesseract.image_to_string(img_array)
            return ocr_text.strip()
        except Exception as e:
            print(f"OCR extraction failed: {e}")
            return ""
    
    def _find_image_caption(self, page, img, page_number: int) -> str:
        """Try to find caption text near an image in PDF."""
        try:
            # This is a simplified caption detection
            # In a real implementation, you might want more sophisticated logic
            text_blocks = page.get_text("dict")
            
            # Look for text blocks that might be captions
            # (This is a basic implementation - could be enhanced)
            captions = []
            for block in text_blocks.get("blocks", []):
                if "lines" in block:
                    for line in block["lines"]:
                        for span in line.get("spans", []):
                            text = span.get("text", "").strip()
                            if text and ("figure" in text.lower() or "image" in text.lower() or "fig" in text.lower()):
                                captions.append(text)
            
            return " ".join(captions) if captions else ""
        except:
            return ""
    
    def _extract_image_context(self, page, img_item, page_text: str, page_number: int, img_index: int) -> dict:
        """使用 PyMuPDF 內建函數精確提取圖片周圍的上下文資訊"""
        try:
            # Step 1: 獲取精確的圖片邊界框
            img_bbox = self._get_precise_image_bbox(page, img_item, img_index)
            
            if img_bbox:
                print(f"   📍 精確 bbox: {img_bbox}")
                
                # Step 2: 使用區域文字提取
                text_first, text_second = self._extract_regional_text(page, img_bbox, img_index)
                
                # Step 3: 組合上下文
                surrounding_text = self._combine_contextual_text(text_first, text_second, img_index)
                
                print(f"   📄 第一部分文字: {text_first[:80]}...")
                print(f"   📄 第二部分文字: {text_second[:80]}...")
                
                # 提取章節標題 - 使用新的 title extractor
                section_title = self.title_extractor._extract_section_title_from_page(page)
                
                # 獲取圖片的所有相關章節
                if img_bbox:
                    img_y_position = img_bbox.y0
                else:
                    img_y_position = None
                
                image_section = self.title_extractor.get_image_section(
                    page_number,  # Use 1-based indexing (no conversion needed)
                    img_y_position
                )
                
                return {
                    'surrounding_text': surrounding_text,
                    'section_title': image_section or "",  # Use enhanced section info or empty string
                    'page_context': page_text[:200] + "..." if len(page_text) > 200 else page_text,
                    'preceding_text': text_first.strip(),  # 可能是左側文字或上方文字
                    'following_text': text_second.strip(),  # 可能是右側文字或下方文字
                    'bbox': img_bbox  # Include the detected bbox
                }
            else:
                print(f"   ⚠️ 無法獲取圖片 {img_index} 的 bbox，使用備選方案")
                return self._fallback_context_extraction(page_text, img_index)
                
        except Exception as e:
            print(f"Error extracting image context: {e}")
            return self._fallback_context_extraction(page_text, img_index)
    
    def _get_precise_image_bbox(self, page, img_item, img_index):
        """使用多種 PyMuPDF 內建方法獲取精確的圖片邊界框"""
        import fitz
        
        # Method 1: 嘗試 page.get_image_bbox (最直接的方法)
        try:
            if hasattr(page, 'get_image_bbox'):
                # 需要使用完整的圖片清單項目
                full_image_list = page.get_images(full=True)
                xref = img_item[0]
                
                # 找到對應的完整圖片項目
                full_img_item = None
                for full_img in full_image_list:
                    if full_img[0] == xref:
                        full_img_item = full_img
                        break
                
                if full_img_item:
                    img_bbox = page.get_image_bbox(full_img_item)
                    if img_bbox and not img_bbox.is_empty:
                        print(f"   ✅ Method 1 成功: get_image_bbox")
                        return img_bbox
        except Exception as e:
            print(f"   Method 1 失敗: {e}")
        
        # Method 2: 使用頁面繪圖對象查找圖片位置
        try:
            xref = img_item[0]
            # 獲取頁面的所有繪圖對象
            drawings = page.get_drawings()
            
            for drawing in drawings:
                # 檢查繪圖對象是否包含我們的圖片
                if hasattr(drawing, 'items'):
                    for item in drawing.items:
                        if hasattr(item, 'type') and item.type == 'i':  # 圖片類型
                            if hasattr(item, 'rect'):
                                img_bbox = fitz.Rect(item.rect)
                                if not img_bbox.is_empty:
                                    print(f"   ✅ Method 2 成功: 從繪圖對象獲取")
                                    return img_bbox
        except Exception as e:
            print(f"   Method 2 失敗: {e}")
        
        # Method 2.5: 使用變換矩陣從完整圖片清單中獲取位置
        try:
            full_image_list = page.get_images(full=True)
            xref = img_item[0]
            
            for full_img in full_image_list:
                if full_img[0] == xref:
                    # 嘗試從完整圖片項目中提取變換矩陣
                    if len(full_img) > 8 and full_img[8]:  # 變換矩陣位置
                        transform = full_img[8]
                        if transform and len(transform) >= 6:
                            # 從變換矩陣計算邊界框
                            a, b, c, d, e, f = transform[:6]
                            # 簡化的邊界框計算
                            x0, y0 = e, f
                            x1, y1 = e + abs(a), f + abs(d)
                            img_bbox = fitz.Rect(x0, y0, x1, y1)
                            if not img_bbox.is_empty:
                                print(f"   ✅ Method 2.5 成功: 從變換矩陣獲取")
                                return img_bbox
                    break
        except Exception as e:
            print(f"   Method 2.5 失敗: {e}")
        
        # Method 3: 使用文字和圖片的相對位置推估
        try:
            # 獲取頁面的所有文字塊
            text_blocks = page.get_text("blocks")
            
            if text_blocks and len(text_blocks) > img_index:
                # 基於文字塊的分布推估圖片位置
                page_height = page.rect.height
                page_width = page.rect.width
                
                # 根據圖片索引分配不同的垂直位置
                if len(text_blocks) >= 4:  # 有足夠多的文字塊
                    # 使用文字塊之間的空隙作為圖片位置
                    gap_positions = []
                    for i in range(len(text_blocks) - 1):
                        current_block = text_blocks[i]
                        next_block = text_blocks[i + 1]
                        gap_y = (current_block[3] + next_block[1]) / 2  # 兩個文字塊之間的中點
                        gap_positions.append(gap_y)
                    
                    if img_index < len(gap_positions):
                        estimated_y = gap_positions[img_index]
                        estimated_bbox = fitz.Rect(50, estimated_y - 25, page_width - 50, estimated_y + 25)
                        print(f"   ✅ Method 3 成功: 基於文字塊間隙估算")
                        return estimated_bbox
            
        except Exception as e:
            print(f"   Method 3 失敗: {e}")
        
        # Method 4: 簡單的等間距分布 (最後備選)
        try:
            page_height = page.rect.height
            page_width = page.rect.width
            
            # 假設有多張圖片時的等間距分布
            total_images = 4  # 假設最多4張圖片
            y_offset = page_height * (0.2 + 0.6 * img_index / max(1, total_images - 1))
            
            estimated_bbox = fitz.Rect(50, y_offset - 30, page_width - 50, y_offset + 30)
            print(f"   ✅ Method 4 成功: 等間距分布估算")
            return estimated_bbox
            
        except Exception as e:
            print(f"   Method 4 失敗: {e}")
        
        return None
    
    def _extract_regional_text(self, page, img_bbox, img_index):
        """基於精確的圖片邊界框提取區域文字，支援內嵌圖片（左右文字）"""
        import fitz
        
        try:
            # 定義搜索半徑
            radius = 50
            
            # 檢查圖片是否為內嵌圖片（小圖片，通常是按鈕或圖標）
            img_width = img_bbox.x1 - img_bbox.x0
            img_height = img_bbox.y1 - img_bbox.y0
            is_inline_image = (img_width < 100 and img_height < 100) or (img_width < 50 or img_height < 50)
            
            if is_inline_image:
                print(f"   🔍 檢測到內嵌圖片 (尺寸: {img_width:.1f}x{img_height:.1f})，使用左右文字提取")
                
                # 對內嵌圖片，提取左側和右側的文字
                # 擴大垂直範圍以包含同一行文字
                line_margin = 10
                
                # 左側區域：從頁面左邊到圖片左邊
                left_rect = fitz.Rect(
                    0,                                    # 左邊界（頁面左邊）
                    max(0, img_bbox.y0),    # 上邊界（圖片上方一點）
                    img_bbox.x0,                          # 右邊界（圖片左邊）
                    min(page.rect.height, img_bbox.y1)  # 下邊界（圖片下方一點）
                )
                
                # 右側區域：從圖片右邊到頁面右邊
                right_rect = fitz.Rect(
                    img_bbox.x1,                          # 左邊界（圖片右邊）
                    max(0, img_bbox.y0),    # 上邊界（圖片上方一點）
                    page.rect.width,                      # 右邊界（頁面右邊）
                    min(page.rect.height, img_bbox.y1)  # 下邊界（圖片下方一點）
                )
                
                text_left = page.get_textbox(left_rect).strip()
                text_right = page.get_textbox(right_rect).strip()
                
                print(f"   🔍 內嵌圖片搜索區域 - 左側: {left_rect}, 右側: {right_rect}")
                print(f"   📝 左側文字: '{text_left[-50:] if len(text_left) > 50 else text_left}'")
                print(f"   📝 右側文字: '{text_right[:50] if len(text_right) > 50 else text_right}'")
                
                # 如果左側或右側文字太少，擴大搜索範圍
                if len(text_left) < 10:
                    expanded_left = fitz.Rect(0, max(0, img_bbox.y0 - 20), img_bbox.x0, min(page.rect.height, img_bbox.y1 + 20))
                    text_left = page.get_textbox(expanded_left).strip()
                    print(f"   📈 擴大左側搜索區域")
                
                if len(text_right) < 10:
                    expanded_right = fitz.Rect(img_bbox.x1, max(0, img_bbox.y0 - 20), page.rect.width, min(page.rect.height, img_bbox.y1 + 20))
                    text_right = page.get_textbox(expanded_right).strip()
                    print(f"   📈 擴大右側搜索區域")
                
                return text_left, text_right
            
            else:
                print(f"   🔍 檢測到塊狀圖片 (尺寸: {img_width:.1f}x{img_height:.1f})，使用上下文字提取")
                
                # 對於大圖片，使用傳統的上下文字提取
                # 上方區域：從頁面頂部到圖片頂部
                above_rect = fitz.Rect(
                    0,                              # 左邊界
                    max(0, img_bbox.y0 - radius),   # 上邊界（圖片上方 + 半徑）
                    page.rect.width,                # 右邊界
                    img_bbox.y0                     # 下邊界（圖片頂部）
                )
                
                # 下方區域：從圖片底部到頁面底部
                below_rect = fitz.Rect(
                    0,                              # 左邊界
                    img_bbox.y1,                    # 上邊界（圖片底部）
                    page.rect.width,                # 右邊界
                    min(page.rect.height, img_bbox.y1 + radius)  # 下邊界（圖片下方 + 半徑）
                )
                
                # 提取區域內的文字
                text_above = page.get_textbox(above_rect).strip()
                text_below = page.get_textbox(below_rect).strip()
                
                print(f"   🔍 塊狀圖片搜索區域 - 上方: {above_rect}, 下方: {below_rect}")
                
                # 如果上方或下方文字太少，擴大搜索範圍
                if len(text_above) < 20:
                    expanded_above = fitz.Rect(0, 0, page.rect.width, img_bbox.y0)
                    text_above = page.get_textbox(expanded_above).strip()[-200:]  # 取最後200字符
                    print(f"   📈 擴大上方搜索區域")
                
                if len(text_below) < 20:
                    expanded_below = fitz.Rect(0, img_bbox.y1, page.rect.width, page.rect.height)  
                    text_below = page.get_textbox(expanded_below).strip()[:200]  # 取前200字符
                    print(f"   📈 擴大下方搜索區域")
                
                return text_above, text_below
            
        except Exception as e:
            print(f"   區域文字提取失敗: {e}")
            # 備選：平分頁面文字
            full_text = page.get_text()
            mid_point = len(full_text) // 2
            return full_text[:mid_point], full_text[mid_point:]
    
    def _combine_contextual_text(self, text_left_or_above, text_right_or_below, img_index):
        """組合上下文文字，添加圖片位置標記，支援左右文字和上下文字"""
        
        # 清理文字
        text_left_or_above = text_left_or_above.strip()
        text_right_or_below = text_right_or_below.strip()
        
        # 檢測是否為左右文字（通常左側文字以介詞或動詞結尾，右側文字以名詞開頭）
        # 或者檢查文字是否在同一行（長度相對較短且沒有句號結尾）
        is_inline_context = (
            len(text_left_or_above) < 100 and len(text_right_or_below) < 100 and
            not text_left_or_above.endswith('.') and not text_left_or_above.endswith('。')
        )
        
        if is_inline_context:
            print(f"   🔤 檢測到內嵌文字上下文（左右結構）")
            # 對於內嵌圖片，保持完整的左右文字
            # 不需要截斷，因為通常都是短文字
            if text_left_or_above and text_right_or_below:
                return f"{text_left_or_above} [IMAGE_HERE] {text_right_or_below}"
            elif text_left_or_above:
                return f"{text_left_or_above} [IMAGE_HERE]"
            elif text_right_or_below:
                return f"[IMAGE_HERE] {text_right_or_below}"
            else:
                return f"[IMAGE_HERE] (圖片 {img_index + 1})"
        else:
            print(f"   🔤 檢測到塊狀文字上下文（上下結構）")
            # 對於塊狀圖片，使用傳統的上下文字處理
            # 如果文字太長，截取合適的長度
            if len(text_left_or_above) > 200:
                # 保留最後200個字符（最接近圖片的部分）
                text_left_or_above = "..." + text_left_or_above[-200:]
            
            if len(text_right_or_below) > 200:
                # 保留前200個字符（最接近圖片的部分）
                text_right_or_below = text_right_or_below[:200] + "..."
            
            # 組合文字，添加圖片位置標記
            if text_left_or_above and text_right_or_below:
                return f"{text_left_or_above} [IMAGE_HERE] {text_right_or_below}"
            elif text_left_or_above:
                return f"{text_left_or_above} [IMAGE_HERE]"
            elif text_right_or_below:
                return f"[IMAGE_HERE] {text_right_or_below}"
            else:
                return f"[IMAGE_HERE] (圖片 {img_index + 1})"
    
    def _extract_section_title_from_page(self, page):
        """從頁面提取章節標題"""
        try:
            # 獲取文字塊並尋找標題樣式的文字
            text_dict = page.get_text("dict")
            
            for block in text_dict.get("blocks", []):
                if "lines" in block:
                    for line in block["lines"]:
                        for span in line.get("spans", []):
                            text = span.get("text", "").strip()
                            size = span.get("size", 0)
                            flags = span.get("flags", 0)
                            
                            # 尋找較大字體或粗體文字作為標題
                            if text and (size > 12 or flags & 16):  # flags & 16 表示粗體
                                if len(text) < 100 and not text.endswith('.'):
                                    return text
            
            return ""
        except:
            return ""
    
    def _fallback_context_extraction(self, page_text, img_index):
        """備選的上下文提取方法"""
        # 按句子分割文字
        sentences = [s.strip() for s in page_text.split('.') if s.strip()]
        
        if not sentences:
            surrounding_text = f"[IMAGE_HERE] (圖片 {img_index + 1})"
            return {
                'surrounding_text': surrounding_text,
                'section_title': '',
                'page_context': page_text[:200],
                'preceding_text': '',
                'following_text': '',
                'bbox': None  # No bbox available in fallback
            }
        
        # 根據圖片索引分配不同的句子
        total_sentences = len(sentences)
        sentences_per_image = max(1, total_sentences // 4)  # 假設最多4張圖片
        
        start_idx = img_index * sentences_per_image
        end_idx = min(total_sentences, (img_index + 1) * sentences_per_image)
        
        preceding_sentences = sentences[max(0, start_idx - 1):start_idx + 1] if start_idx > 0 else sentences[:1]
        following_sentences = sentences[start_idx + 1:end_idx + 1] if end_idx < total_sentences else sentences[-1:]
        
        preceding_text = '. '.join(preceding_sentences)
        following_text = '. '.join(following_sentences)
        
        surrounding_text = f"{preceding_text} [IMAGE_HERE] {following_text}"
        
        return {
            'surrounding_text': surrounding_text,
            'section_title': '',
            'page_context': page_text[:200],
            'preceding_text': preceding_text,
            'following_text': following_text,
            'bbox': None  # No bbox available in fallback
        }
    
    def _extract_section_title_fallback(self, page_text: str) -> str:
        """Fallback method to extract section title from page text."""
        import re
        # Look for common section title patterns
        lines = page_text.split('\n')
        for line in lines[:10]:  # Check first 10 lines
            line = line.strip()
            if line and len(line) < 100:  # Likely a title if short
                # Check if it looks like a section title
                if re.match(r'^[0-9]+\..*', line) or line.isupper() or len(line.split()) <= 8:
                    return line
        return ''
    
    def _extract_section_title(self, text_blocks) -> str:
        """提取章節標題（基於字體大小判斷）"""
        try:
            titles = []
            max_font_size = 0
            
            for block in text_blocks.get("blocks", []):
                if "lines" in block:
                    for line in block["lines"]:
                        for span in line.get("spans", []):
                            text = span.get("text", "").strip()
                            font_size = span.get("size", 0)
                            
                            if text and font_size > max_font_size:
                                max_font_size = font_size
                                titles = [text]
                            elif text and font_size == max_font_size:
                                titles.append(text)
            
            # 返回最大字體的文字作為標題
            return " ".join(titles[:2]) if titles else ""
            
        except:
            return ""
    
    def _extract_docx_section_title(self, doc: DocumentType) -> str:
        """從DOCX文件提取章節標題（基於樣式判斷）"""
        try:
            # Look for headings in the document
            for paragraph in doc.paragraphs[:10]:  # Check first 10 paragraphs
                if paragraph.style.name.startswith('Heading') or 'Title' in paragraph.style.name:
                    text = paragraph.text.strip()
                    if text and len(text) < 100:  # Reasonable title length
                        return text
            
            # Fallback: look for first non-empty paragraph that might be a title
            for paragraph in doc.paragraphs[:5]:
                text = paragraph.text.strip()
                if text and len(text) < 100:
                    # Check if it looks like a title (short, not sentence-like)
                    if not text.endswith('.') and len(text.split()) <= 10:
                        return text
            
            return ""
        except Exception as e:
            print(f"Error extracting DOCX section title: {e}")
            return ""
    
    def _extract_txt_section_title(self, text: str) -> str:
        """從TXT文件提取章節標題（基於首行判斷）"""
        try:
            lines = text.split('\n')
            for line in lines[:5]:  # Check first 5 lines
                line = line.strip()
                if line and len(line) < 100:  # Reasonable title length
                    # Check if it looks like a title
                    if not line.endswith('.') and len(line.split()) <= 15:
                        return line
            return ""
        except Exception as e:
            print(f"Error extracting TXT section title: {e}")
            return ""
    
    def _extract_table_text(self, table: Table) -> str:
        """Extract text from a DOCX table."""
        table_text = []
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                table_text.append(" | ".join(row_text))
        
        return "\n".join(table_text)
    
    def save_image(self, image_data: ImageData, save_dir: str) -> str:
        """Save image to disk and return the file path."""
        os.makedirs(save_dir, exist_ok=True)
        
        file_path = os.path.join(save_dir, f"{image_data.image_id}.png")
        image_data.image.save(file_path, "PNG")
        
        return file_path
    
    def image_to_base64(self, image: Image.Image) -> str:
        """Convert PIL Image to base64 string for web display."""
        buffer = io.BytesIO()
        image.save(buffer, format="PNG")
        img_str = base64.b64encode(buffer.getvalue()).decode()
        return f"data:image/png;base64,{img_str}"